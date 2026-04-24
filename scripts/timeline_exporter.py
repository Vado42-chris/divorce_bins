import os
import sys
import json
import logging
from datetime import datetime

# Optional graceful degradation if python-docx isn't installed natively
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx is not installed. Exporting to structured Markdown instead.")

METADATA_FILE = "/media/chrishallberg/Storage 22/002_Personal/divorce_bins/metadata/index.json"
EXPORT_DIR = "/media/chrishallberg/Storage 22/002_Personal/divorce_bins/exports"

def generate_chronological_timeline(case_name="Hallberg v. Andersen", scope="Global"):
    os.makedirs(EXPORT_DIR, exist_ok=True)
    
    with open(METADATA_FILE, 'r') as f:
        master_index = json.load(f)

    # Filter to specific domains or global overlap
    if scope != "Global":
        master_index = [m for m in master_index if m.get("domain") == scope]

    # Chronological sort (oldest to newest for court timeline sequence)
    master_index.sort(key=lambda x: x.get('timestamp', '0000-00-00'))

    if DOCX_AVAILABLE:
        doc = Document()
        # Title Page / Header
        title = doc.add_heading(f'EVIDENCE TIMELINE & AFFIDAVIT SUPPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Matter: {case_name}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Scope Limit: {scope}")
        doc.add_page_break()

        doc.add_heading('Chronological Exhibit Ledger', level=1)
        
        for item in master_index:
            date_str = item.get('timestamp', 'Unknown Date')
            p = doc.add_paragraph()
            p.add_run(f"[{date_str}] - ").bold = True
            p.add_run(f"{item.get('title', 'Unknown Record')}")
            
            # SGI Structured Injection
            if 'sgi_metadata' in item:
                sgi = item['sgi_metadata']
                meta_p = doc.add_paragraph(style='List Bullet')
                meta_p.add_run(f"SGI Claim: {sgi.get('claimNumber')} | Severity: {sgi.get('injurySeverity')}").italic = True
            
            # Semantic summaries if intelligence is active
            if 'summary' in item:
                doc.add_paragraph(f"> {item['summary']}")
        
        out_path = os.path.join(EXPORT_DIR, f"Timeline_{case_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx")
        doc.save(out_path)
        print(f"Generated strict Legal DOCX Exhibit Ledger: {out_path}")
        return out_path
    else:
        # Fallback to pure highly-structured markdown if docx is missing in venv
        out_path = os.path.join(EXPORT_DIR, f"Timeline_{case_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.md")
        with open(out_path, 'w') as f:
            f.write(f"# EVIDENCE TIMELINE: {case_name}\n")
            f.write(f"Generated: {datetime.now().strftime('%B %d, %Y')}\n\n")
            for item in master_index:
                f.write(f"### {item.get('timestamp')} - {item.get('title')}\n")
                if 'sgi_metadata' in item:
                    f.write(f"**SGI Match:** Claim {item['sgi_metadata'].get('claimNumber')}\n")
                f.write("\n")
        print(f"Generated Markdown Ledger (python-docx not found): {out_path}")
        return out_path

if __name__ == "__main__":
    scope_arg = sys.argv[1] if len(sys.argv) > 1 else 'Global'
    generate_chronological_timeline(scope=scope_arg)
